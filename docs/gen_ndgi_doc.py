#!/usr/bin/env python3
"""Generate NDGI Technical Documentation in DOCX format."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Consolas'
font.size = Pt(10.5)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_code(text):
    p = doc.add_paragraph()
    for line in text.strip().split('\n'):
        if line == '':
            run = p.add_run(' ')
        else:
            run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if line != text.strip().split('\n')[-1]:
            run = p.add_run('\n')
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')

# ===== Title =====
title = doc.add_heading('NDGI: Neural Dynamic Global Illumination', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_para(f'Version 1.0 | {datetime.date.today().strftime("%Y-%m-%d")} | SomeGI Engine')
add_para('Authors: NDGI Implementation Team')
doc.add_page_break()

# ===== 1 =====
add_heading('1. Overview', level=1)
add_para('NDGI (Neural Dynamic Global Illumination) is a real-time global illumination technique '
         'based on a miniature neural network. It uses a tiny MLP (Multi-Layer Perceptron) as an '
         '"irradiance cache" to replace traditional spherical harmonic probe grids (e.g., DDGI). '
         'NDGI continuously learns the lighting distribution in the scene through online training, '
         'supporting dynamic light sources and dynamic scenes.')
add_para('The core concept is derived from Muller et al. 2021 (SIGGRAPH) "Neural Radiance Caching", '
         'but significantly simplified for non-Tensor-Core GPUs (e.g., Intel UHD 770): '
         'the 8-layer x 64-neuron MLP is compressed to 2 layers x 16 neurons, using generic '
         'compute shaders for both inference and training.')

add_heading('1.1 Key Features', level=2)
add_bullet('Tiny MLP: 2 hidden layers x 16 neurons, 451 parameters total (~1.8 KB)')
add_bullet('Online training: Collects samples from probe ray tracing each frame, updates weights via mini-batch SGD')
add_bullet('EMA smoothing: Exponential moving average (alpha=0.95) reduces frame-to-frame flickering')
add_bullet('Universal GPU: No Tensor Cores required, pure compute shader implementation')
add_bullet('Decoupled from rendering pipeline: Supports both forward and deferred rendering modes')
doc.add_page_break()

# ===== 2 =====
add_heading('2. MLP Architecture', level=1)
add_para('The NDGI neural network is a three-layer fully-connected network:')
add_code('Input(6) -> Hidden1(16) -> Hidden2(16) -> Output(3)\nActivation: Leaky ReLU (alpha=0.01), Output: ReLU (irradiance >= 0)')

add_heading('2.1 Input Features', level=2)
add_para('The network receives a 6-dimensional input vector encoding geometry at a surface point:')
add_code('Input = [worldPos.x, worldPos.y, worldPos.z, normal.x, normal.y, normal.z]')
add_para('These 6 values allow the network to learn "outgoing irradiance at position P with normal N". '
         'Compared to Muller 2021, albedo and depth are omitted (irradiance is material-independent).')

add_heading('2.2 Weight Parameters', level=2)
add_para('Weights are stored in float4-packed format in Storage Buffers:')
add_code('Layer 1: W1[24 float4] = 6x16 = 96 floats;  B1[4 float4] = 16 floats\n'
         'Layer 2: W2[64 float4] = 16x16 = 256 floats; B2[4 float4] = 16 floats\n'
         'Layer 3: W3[12 float4] = 16x3 = 48 floats;  B3[1 float4] = 3 floats\n'
         'Total: 451 floats ~ 1.8 KB (6 Storage Buffers)')

add_heading('2.3 Forward Pass (Inference)', level=2)
add_para('Mathematical formulation of a single inference:')
add_code('h1 = LeakyReLU(W1 * [pos, N] + B1)    // 6 -> 16\n'
         'h2 = LeakyReLU(W2 * h1 + B2)            // 16 -> 16\n'
         'irradiance = ReLU(W3 * h2 + B3)         // 16 -> 3 (RGB)')

add_heading('2.4 Float4 Packing', level=2)
add_para('To simplify Storage Buffer declarations, weights are packed into float4 arrays. '
         'For W1 (96 floats, 24 float4): W1[0] = weights from input[0] to neurons (0,1,2,3). '
         'This packing scheme applies to all weight matrices with modulo-4 alignment.')
doc.add_page_break()

# ===== 3 =====
add_heading('3. Inference Process', level=1)
add_para('Inference occurs in the lighting shader (lighting.slang / forward_ibl.slang), '
         'executed once per pixel.')

add_heading('3.1 Deferred Rendering Path', level=2)
add_para('In lighting.slang compute shader:')
add_code('// ddgiCounts.w == 2 triggers NDGI path\n'
         'float3 worldPos = worldFromDepth(pix, depth);\n'
         'float3 N = normalize(gNormalRough.Load(pix).xyz);\n'
         'float3 ndgiIrr = ndgiSampleIrradiance(worldPos, N);\n'
         'diffuse = ndgiIrr * albedo * (1 - metallic) / PI;')

add_heading('3.2 Forward Rendering Path', level=2)
add_para('In forward_ibl.slang fragment shader:')
add_code('float3 N = normalize(i.normal);\n'
         'float3 ndgiIrr = ndgiSampleIrradiance(i.worldPos, N);\n'
         'indirect = ndgiIrr * base.rgb * (1.0 - metallic) / PI;')

add_heading('3.3 Inference Performance', level=2)
add_para('Per-inference compute cost (in float MAD operations):')
add_code('Layer 1: 6 x 16 = 96 MAD + 16 bias\n'
         'Layer 2: 16 x 16 = 256 MAD + 16 bias\n'
         'Layer 3: 16 x 3 = 48 MAD + 3 bias\n'
         'Total: ~400 MAD / pixel (+ activation functions)')
add_para('At 1920x1080, only geometry-covered pixels execute (typically <50%), '
         'resulting in ~400M MAD/frame. Estimated <2ms on Intel UHD 770.')
doc.add_page_break()

# ===== 4 =====
add_heading('4. Training Process', level=1)

add_heading('4.1 Training Data Collection', level=2)
add_para('Each frame, samples are collected by tracing rays from 256 probes (8x4x8 grid):')
add_code('32 rays per probe (spherical Fibonacci sampling + per-frame rotation)\n'
         'Total: 256 x 32 = 8,192 training samples per frame\n\n'
         'Sample format (9 floats):\n'
         '  [worldPos.xyz, normal.xyz, targetRadiance.rgb]\n\n'
         'Target computation:\n'
         '  RayQuery -> hit surface -> directBRDF(sun) + emissive + ambient')
add_para('Uses VK_KHR_ray_query for hardware-accelerated ray tracing. '
         'For each ray from each probe, a ray is cast from the probe position, '
         'traced to the hit surface, and direct lighting (sun + emissive + ambient) '
         'is computed as the training target.')

add_heading('4.2 Training Algorithm', level=2)
add_para('Training runs in a dedicated compute shader (ndgi_train.slang):')
add_code('Per-frame training loop:\n'
         '  1. Read sampleCount from probe trace output\n'
         '  2. For iteration in [1..4]:\n'
         '     a. Randomly select 256 samples (mini-batch)\n'
         '     b. Forward pass: predict radiance for each sample\n'
         '     c. Compute MSE loss: L = mean((pred - target)^2)\n'
         '     d. Backward pass: compute gradients via backpropagation\n'
         '     e. SGD update: W -= lr * dL/dW\n'
         '  3. EMA smooth: W = alpha * W_old + (1-alpha) * W_new  (alpha=0.95)\n\n'
         'Hyperparameters:\n'
         '  learning_rate  = 0.01\n'
         '  batch_size     = 256\n'
         '  num_iterations = 4 per frame\n'
         '  ema_alpha      = 0.95')

add_heading('4.3 Backpropagation', level=2)
add_para('Gradients are computed via the chain rule for each layer:')
add_code('Layer 3 (16->3):\n'
         '  dL/dW3 = outGrad outerProd h2\n'
         '  dL/dB3 = outGrad\n'
         '  dL/dh2 = W3^T * outGrad\n\n'
         'Layer 2 (16->16):\n'
         '  delta = dL/dh2 * LeakyReLU_deriv(h2)\n'
         '  dL/dW2 = delta outerProd h1\n'
         '  dL/dB2 = delta\n'
         '  dL/dh1 = W2^T * delta\n\n'
         'Layer 1 (6->16):\n'
         '  delta = dL/dh1 * LeakyReLU_deriv(h1)\n'
         '  dL/dW1 = delta outerProd [pos, N]\n'
         '  dL/dB1 = delta')

add_heading('4.4 EMA Smoothing', level=2)
add_para('Exponential Moving Average reduces weight jitter across frames:')
add_code('W_ema = alpha * W_previous + (1-alpha) * W_sgd')
add_para('Alpha=0.95 means new gradients contribute only 5% to the weight update. '
         'This high alpha value serves multiple purposes:')
add_bullet('Smooths frame-to-frame lighting changes, reducing flicker')
add_bullet('Filters single-frame sample noise (only 8,192 samples/frame)')
add_bullet('Allows lighting to gradually converge to a stable state over multiple frames')
doc.add_page_break()

# ===== 5 =====
add_heading('5. Probe Ray Tracing', level=1)
add_para('NDGI reuses the DDGI probe grid system for spatial coverage.')

add_heading('5.1 Probe Layout', level=2)
add_code('Grid: 8 x 4 x 8 = 256 probes\n'
         'Spacing: auto-computed from scene AABB\n'
         '  spacing = max(aabb_size) * 1.05 / (probes_per_dim - 1)')
add_para('Probes are uniformly distributed in the scene bounding box, covering all reachable areas.')

add_heading('5.2 Ray Direction Sampling', level=2)
add_para('Ray directions use the Spherical Fibonacci sequence for near-uniform sphere coverage:')
add_code('phi = 2*PI * frac(i * (phi_golden - 1))   // golden ratio = 1.618...\n'
         'cosTheta = 1 - (2*i + 1) / N\n'
         'sinTheta = sqrt(1 - cosTheta^2)\n'
         'dir = (cos(phi)*sinTheta, sin(phi)*sinTheta, cosTheta)')
add_para('The Fibonacci sequence produces near-uniform distribution on the sphere with '
         'better coverage than random sampling. A per-frame Y-axis rotation (jitter) '
         'avoids directional bias over time.')

add_heading('5.3 Ray Query', level=2)
add_para('Hardware-accelerated ray tracing via VK_KHR_ray_query in compute shader:')
add_code('RayQuery<CULL_NON_OPAQUE> q;\n'
         'q.TraceRayInline(gTLAS, RAY_FLAG_NONE, 0xFF, ray);\n'
         'q.Proceed();\n'
         'if (q.CommittedStatus() == TRIANGLE_HIT) {\n'
         '    // Read hit attributes: position, normal, UV, material\n'
         '    // Compute direct lighting as training target\n'
         '}')
doc.add_page_break()

# ===== 6 =====
add_heading('6. System Architecture', level=1)

add_heading('6.1 File Structure', level=2)
add_code('src/renderer/ndgi_resources.h/cpp       -- MLP weights + sample buffer management\n'
         'src/renderer/ndgi_pass.h/cpp            -- Probe trace + training compute pipelines\n'
         'shaders/gi/ndgi/ndgi_infer.slang        -- MLP inference module (imported by lighting)\n'
         'shaders/gi/ndgi/ndgi_init.slang          -- Xavier weight initialization\n'
         'shaders/gi/ndgi/ndgi_train.slang         -- Online training (backprop + SGD + EMA)\n'
         'shaders/gi/ndgi/ndgi_probe_trace.slang   -- Probe ray tracing for sample collection\n'
         'shaders/lighting/lighting.slang          -- Deferred lighting (imports ndgi_infer)\n'
         'shaders/forward/forward_ibl.slang        -- Forward lighting (inline NDGI inference)')

add_heading('6.2 Per-Frame Execution Flow', level=2)
add_code('Pipeline Step "NDGI" (before TS-GI, before Lighting):\n'
         '  1. [First Frame Only] initWeights() -> Xavier random init MLP\n'
         '  2. record() -> ndgi_probe_trace.slang\n'
         '     dispatch(128, 1, 1): 8,192 threads, 1 ray per thread\n'
         '     atomicAdd -> write sample buffer + sample count\n'
         '  3. recordTraining() -> ndgi_train.slang\n'
         '     dispatch(1, 1, 1): single thread, 4 epochs of SGD\n'
         '     unpack weights -> forward -> backward -> SGD -> EMA -> pack weights\n\n'
         'Pipeline Step "Lighting" or ForwardPass::record:\n'
         '  4. Each pixel calls ndgiSampleIrradiance(worldPos, normal)\n'
         '     6->16->16->3 MLP forward pass -> predicted irradiance')

add_heading('6.3 Descriptor Bindings', level=2)
add_para('NDGI weights are passed to shaders via descriptor set=0:')
add_code('Deferred (lighting.slang):   bindings 27-32\n'
         'Forward (forward_ibl.slang): bindings 4-9\n\n'
         'Binding contents:\n'
         '  27/4:  gNdgiW1 (24 float4)  -- Layer 1 weights (96 floats)\n'
         '  28/5:  gNdgiB1 (4 float4)   -- Layer 1 biases  (16 floats)\n'
         '  29/6:  gNdgiW2 (64 float4)  -- Layer 2 weights (256 floats)\n'
         '  30/7:  gNdgiB2 (4 float4)   -- Layer 2 biases  (16 floats)\n'
         '  31/8:  gNdgiW3 (12 float4)  -- Layer 3 weights (48 floats)\n'
         '  32/9:  gNdgiB3 (1 float4)   -- Layer 3 biases  (3 floats)')

add_heading('6.4 Descriptor Set Layout', level=2)
add_para('The training compute pipeline uses 8 bindings (all STORAGE_BUFFER):')
add_code('Bindings 0-5: W1, B1, W2, B2, W3, B3 (RW, for training updates)\n'
         'Binding 6:    sample buffer (read-only, probe trace output)\n'
         'Binding 7:    sample count (read-only, atomic counter)')
doc.add_page_break()

# ===== 7 =====
add_heading('7. Theoretical Foundations', level=1)

add_heading('7.1 Neural Radiance Caching', level=2)
add_para('Traditional GI methods (e.g., DDGI) use fixed-resolution probe grids with '
         'Spherical Harmonics (SH) to store irradiance. The main limitations of SH are: '
         'spatial resolution limited by probe density, angular resolution limited by SH order, '
         'and complex interpolation logic (trilinear interpolation + backface culling + Chebyshev visibility).')
add_para('Neural networks as "universal function approximators" can represent the scene irradiance field '
         'more compactly:')
add_bullet('Continuous representation: Unlike discrete probes, MLP predicts at any point in space')
add_bullet('Adaptive resolution: Network automatically allocates more "capacity" where lighting varies sharply')
add_bullet('Compression efficiency: 451 parameters ~ 1.8 KB vs DDGI 64x256x16B = 256 KB irradiance atlas')
add_bullet('No interpolation artifacts: MLP inference produces naturally smooth output')

add_heading('7.2 Self-Training Principle', level=2)
add_para('NDGI uses self-training to implicitly propagate multi-bounce lighting from '
         'single-bounce ray tracing data:')
add_code('Training target = DirectLighting(hit_point) + MLP_prev(hit_point)')
add_para('Although training data only contains direct lighting (sun + emissive), the MLP learns '
         'to correlate irradiance across different points in space during training. As more frames '
         'accumulate, indirect lighting gradually propagates through the scene, equivalent to '
         'multiple bounces - but the computation cost per ray remains constant.')
add_para('This is conceptually similar to Muller 2021 "training suffix", but simplified: '
         'instead of tracing additional path bounces, we rely on the MLP self-prediction '
         'across frames as a proxy for indirect lighting.')

add_heading('7.3 Comparison with DDGI', level=2)
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Dimension'
hdr[1].text = 'DDGI'
hdr[2].text = 'NDGI'
rows_data = [
    ['Storage', 'Probe atlas (256 KB+)', 'MLP weights (1.8 KB)'],
    ['Query', 'Trilinear interpolation + visibility check', 'MLP forward pass (400 MAD)'],
    ['Training', 'None (EMA updates texels directly)', 'SGD + backprop (per-frame)'],
    ['Spatial Repr.', 'Discrete probe grid', 'Continuous function (MLP)'],
    ['Temporal Smooth.', 'Texel EMA (0.92)', 'Weight EMA (0.95)'],
    ['Hardware', 'Any GPU', 'Any GPU + RT (ray query)'],
    ['New Scene', 'Instant', 'Requires seconds of training to converge'],
]
for i, row_data in enumerate(rows_data):
    row = table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
doc.add_page_break()

# ===== 8 =====
add_heading('8. Implementation Details', level=1)

add_heading('8.1 Weight Initialization', level=2)
add_para('Weights are initialized using Xavier (Glorot) uniform distribution:')
add_code('W ~ Uniform(-sqrt(6/(fan_in+fan_out)), +sqrt(6/(fan_in+fan_out)))\n'
         'B = 0\n\n'
         'Specific scales:\n'
         '  Layer 1 (6->16): scale = sqrt(6/22) ~ 0.522\n'
         '  Layer 2 (16->16): scale = sqrt(6/32) ~ 0.433\n'
         '  Layer 3 (16->3): scale = sqrt(6/19) ~ 0.562')
add_para('Xavier initialization ensures that the variance of activations remains stable '
         'across layers, preventing vanishing/exploding gradients at training start.')

add_heading('8.2 Buffer Management', level=2)
add_para('Weight buffers are created as device-local storage buffers with '
         'VK_BUFFER_USAGE_STORAGE_BUFFER_BIT. The sample buffer (288 KB for 8,192 x 9 floats) '
         'is also device-local for GPU-only access. The sample count buffer is host-visible '
         'for CPU readback during training dispatch setup.')

add_heading('8.3 Shader Compilation', level=2)
add_para('NDGI shaders are compiled with Slang to SPIR-V targeting spirv_1_5. '
         'The ndgi_infer.slang module is imported by lighting.slang. '
         'The forward_ibl.slang duplicates the inference function inline '
         'to avoid binding number conflicts (bindings 4-9 vs 27-32).')

add_heading('8.4 Pixel Loop Integration', level=2)
add_para('In deferred mode, the NDGI inference is inserted into the pixel loop '
         'of lighting.slang\'s compute shader at the point where DDGI samples are normally taken. '
         'The ddgiCounts.w flag (value 2) selects NDGI over DDGI (value 1) or none (value 0). '
         'This flag is set in App::run() based on m_ndgiEnabled.')

doc.add_page_break()

# ===== 9 =====
add_heading('9. References', level=1)
add_bullet('Muller, T., Rousselle, F., Novak, J., & Keller, A. (2021). '
           '"Real-time Neural Radiance Caching for Path Tracing." ACM Trans. Graph. 40(4), Article 69.')
add_bullet('Majercik, Z., Guertin, J.-P., Nowrouzezahrai, D., & McGuire, M. (2019). '
           '"Dynamic Diffuse Global Illumination with Ray-Traced Irradiance Fields." JCGT 8(2).')
add_bullet('NVIDIA RTXGI SDK. https://github.com/NVIDIA-RTX/RTXGI')
add_bullet('Wu, J., Zhou, J., Zhou, Z., Huang, Z., & Li, C. (2026). '
           '"Neural Dynamic GI: Random-Access Neural Compression for Temporal Lightmaps." CVPR 2026.')

# Save
os.makedirs('D:/Source/SomeGI/docs', exist_ok=True)
output_path = 'D:/Source/SomeGI/docs/NDGI_Technical_Documentation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Size: {os.path.getsize(output_path)} bytes')
