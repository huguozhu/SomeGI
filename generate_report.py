#!/usr/bin/env python3
"""生成 SomeGI 全局光照技术实现报告 (docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    hfont = heading_style.font
    hfont.name = '微软雅黑'
    hfont.color.rgb = RGBColor(0x1A, 0x56, 0xDB) if level <= 2 else RGBColor(0x33, 0x33, 0x33)
    hfont.bold = True
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_paragraph(text, style_name='Normal', bold=False, color=None, size=None, alignment=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if bold: run.bold = True
    if color: run.font.color.rgb = color
    if size: run.font.size = size
    if alignment is not None: p.alignment = alignment
    return p

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '1A56DB')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ── Cover Page ──────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SomeGI 全局光照技术实现报告')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Vulkan 实时全局光照实验平台 — 技术原理、实现细节与性能对比')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('版本: 2026-05-19 | 平台: Windows 11 / Vulkan 1.3 / MSVC 2026 | GPU: Intel UHD 770')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ── 目录占位 ────────────────────────────────────────────
doc.add_heading('目录', level=1)
toc_items = [
    '1. 项目概述',
    '2. 渲染管线架构',
    '3. IBL — 基于图像的光照 (Image-Based Lighting)',
    '4. 屏幕空间技术：SSAO / GTAO / SSR / SSGI / GTGI',
    '5. RSM — 反射阴影贴图 (Reflective Shadow Maps)',
    '6. LPV — 光传播体积 (Light Propagation Volumes)',
    '7. VXGI — 体素锥追踪全局光照 (Voxel Cone Tracing GI)',
    '8. PRT — 预计算辐射度传输 (Precomputed Radiance Transfer)',
    '9. DDGI — 动态漫反射全局光照 (Dynamic Diffuse GI)',
    '10. SDFGI — 有符号距离场全局光照 (Signed Distance Field GI)',
    '11. ReSTIR DI — 储层时空重要性重采样直接光照',
    '12. 性能与效果对比',
    '13. 总结与展望',
]
for item in toc_items:
    doc.add_paragraph(item, style='Normal')

doc.add_page_break()

# ── 1. 项目概述 ─────────────────────────────────────────
doc.add_heading('1. 项目概述', level=1)

doc.add_paragraph(
    'SomeGI 是一个基于 C++20 / Vulkan 1.3 的实时全局光照（Global Illumination, GI）实验平台。'
    '项目以可切换方式集成了十种以上主流的实时 GI 技术，涵盖从轻量级屏幕空间方法到完整的体积光传输方案，'
    '并共享统一的延迟渲染管线、场景加载器和材质系统。通过 ImGui 下拉框即可实时切换不同的 GI 技术，'
    '为算法对比、技术选型和教学演示提供了便利。'
)

doc.add_heading('技术栈', level=2)
add_table(
    ['组件', '技术选型'],
    [
        ['语言', 'C++20'],
        ['图形 API', 'Vulkan 1.3'],
        ['Shader 语言', 'Slang → SPIR-V'],
        ['窗口库', 'GLFW 3.4'],
        ['数学库', 'GLM'],
        ['设备初始化', 'vk-bootstrap'],
        ['场景加载', 'cgltf (glTF 2.0)'],
        ['UI', 'Dear ImGui (Docking 分支)'],
        ['构建系统', 'CMake 3.24+ / MSVC 2026'],
        ['测试模型', 'Khronos glTF-Sample-Models (Cube, Sponza)'],
    ],
    [4, 10]
)

doc.add_heading('已实现的 GI 技术一览', level=2)
add_table(
    ['技术', '缩写', '类型', '阶段', '简介'],
    [
        ['基于图像的光照', 'IBL', '环境光探针', 'M3', 'Split-sum 近似；预过滤环境贴图 + BRDF LUT'],
        ['屏幕空间环境光遮蔽', 'SSAO', '屏幕空间', 'M4.1', '逐像素半球采样，计算局部遮挡'],
        ['Ground Truth AO', 'GTAO', '屏幕空间', 'M4.1', '基于水平线的解析积分，替换 SSAO'],
        ['屏幕空间反射', 'SSR', '屏幕空间', 'M4.2', '屏幕空间 ray march，镜面反射颜色'],
        ['屏幕空间 GI', 'SSGI', '屏幕空间', 'M4.3', '半球 ray march 采样上一帧 HDR，一次反弹漫反射'],
        ['Ground Truth GI', 'GTGI', '屏幕空间', 'C.5', '切片水平线搜索 + 间接辐射度累积，替代 SSGI'],
        ['反射阴影贴图', 'RSM', '体积/空间', 'M5', '太阳视角 G-Buffer + VPL gather'],
        ['光传播体积', 'LPV', '体积/空间', 'M6', '32³ SH 网格注入 + 6 邻居迭代传播'],
        ['体素锥追踪 GI', 'VXGI', '体积/空间', 'M7', '128³ 体素化 + mip 链 + 6-cone 追踪'],
        ['预计算辐射度传输', 'PRT', '预计算', 'M8', '32³ 体素可见性 SH 烘焙 + 运行时重建'],
        ['动态漫反射 GI', 'DDGI', '探针/空间', 'M11', '256 探针 + 八面体 atlas + Chebyshev 可见性'],
        ['有符号距离场 GI', 'SDFGI', '体积/空间', 'C.3', 'JFA 构建 UDF + sphere tracing'],
        ['储层时空重采样', 'ReSTIR DI', '采样框架', 'C.4', 'RIS 候选采样 + 空间复用 + 单次可见性'],
    ],
    [3, 2.5, 2.5, 2, 6]
)

doc.add_page_break()

# ── 2. 渲染管线架构 ─────────────────────────────────────
doc.add_heading('2. 渲染管线架构', level=1)

doc.add_paragraph(
    'SomeGI 采用延迟渲染（Deferred Rendering）管线，将场景几何与光照计算分离。'
    '渲染过程分为以下阶段：'
)

doc.add_heading('2.1 管线阶段', level=2)

pipeline_stages = [
    ['阶段 0: RSM 几何', '从太阳视角渲染场景 → 4 张 MRT（位置/法线/通量/深度）'],
    ['阶段 1: G-Buffer', '从相机视角渲染场景 → 3 张 MRT（Albedo+Metallic / Normal+Roughness / Emissive+AO）+ 深度'],
    ['阶段 1.5: AO', 'SSAO 或 GTAO 计算屏幕空间环境光遮蔽 → R8 纹理'],
    ['阶段 1.6: SSR', '屏幕空间 ray march 计算镜面反射 → RGBA16F 纹理'],
    ['阶段 1.7: SSGI/GTGI', '屏幕空间间接漫反射 → RGBA16F 纹理'],
    ['阶段 1.83: VXGI', '体素化 → 光照注入 → mipmap → 可选各向异性/relight'],
    ['阶段 1.835: SDFGI', 'JFA 构建 UDF → sphere tracing → 间接光照'],
    ['阶段 1.836: ReSTIR', 'RIS 初始化 → 空间复用 → 着色（含可见性测试）'],
    ['阶段 1.84: DDGI', '探针光线追踪 → atlas 混合 → 探针分类'],
    ['阶段 1.85: LPV', 'RSM 注入 SH 网格 → 多次 6-邻居传播'],
    ['阶段 1.8: RSM 采样', '屏幕像素 gather RSM VPL → 间接光照'],
    ['阶段 2: 延迟光照', 'Compute shader 整合所有 GI 贡献 → HDR 输出'],
    ['阶段 3: 色调映射', 'HDR → LDR + 伽马校正'],
    ['阶段 4: ImGui', '渲染 UI 叠加层'],
]
for stage, desc in pipeline_stages:
    p = doc.add_paragraph()
    run = p.add_run(stage + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.2 光照计算统一入口', level=2)
doc.add_paragraph(
    '所有 GI 技术的输出通过统一的延迟光照 Compute Shader（lighting.slang）进行整合。'
    '该 Shader 包含 25 个 set=0 绑定（G-Buffer、各 GI 输出纹理/3D 纹理）和 5 个 set=1 绑定（IBL 资源），'
    '通过 FrameUniforms 中的计数/标志字段决定每帧启用哪些 GI 路径。'
    '漫反射和镜面反射分开处理，各 GI 技术按优先级顺序覆盖或混合漫反射项。'
)

doc.add_page_break()

# ── 3-11. 各 GI 技术详细说明 ────────────────────────────

gi_techniques = [
    {
        'num': '3',
        'title': 'IBL — 基于图像的光照 (Image-Based Lighting)',
        'subtitle': 'Split-Sum 近似 (Karis 2013)',
        'intro': (
            'IBL 是 SomeGI 的 GI 基础设施和默认回退方案。它使用一张 HDR 等距柱状环境贴图来近似场景周围的环境光照，'
            '通过预计算将昂贵的半球积分拆分为两个独立的二维查找表，在运行时仅需 3 次纹理采样即可得到完整的 PBR 间接光照。'
        ),
        'principle': (
            '渲染方程的反射项是一个五维积分：方向（2D）、粗糙度（1D）、NoV（1D）、反照率（1D）。'
            'Karis 2013 提出的 Split-Sum 近似将其拆分为两项独立的乘积：\n\n'
            '  L_o ≈ (∫ L_i · D · cosθ dω) · (∫ BRDF · cosθ dω)\n\n'
            '第一项是 GGX 分布加权的预过滤环境贴图（Prefiltered Environment Map），存储为带 mip 链的立方体贴图，'
            '每个 mip 级别对应不同的粗糙度值。运行时根据表面粗糙度选择相应的 mip，粗糙度越高采样越模糊。\n\n'
            '第二项是 BRDF 积分查找表（BRDF LUT），是一个 256×256 的 RG16F 2D 纹理，'
            '以 (NoV, roughness) 为坐标存储两个系数 (A, B)，运行时重建镜面反射：'
            'specular = prefilteredColor · (F0 · A + B)。\n\n'
            '漫反射项更简单：用一个低分辨率（32²）的辐照度立方体贴图（余弦核卷积）表达，沿法线方向采样一次。'
        ),
        'implementation': (
            '烘焙管线（IblBaker）：\n'
            '1. 加载 HDR equirectangular → 上传到 GPU\n'
            '2. equi_to_cube: 等距柱状图 → 512²×6 面 cubemap\n'
            '3. Blit 生成 cubemap mip 链\n'
            '4. prefilter_diffuse: 黎曼求和 → 32² 辐照度 cubemap\n'
            '5. prefilter_specular: GGX 重要性采样 1024 样本/mip → 256² 多 mip cubemap\n'
            '6. brdf_lut: Hammersley 采样 → 256² RG16F LUT\n\n'
            '资源规格：\n'
            '- envCube: 512×512×6, 6 mips, RGBA16F (~10 MB)\n'
            '- diffuseCube: 32×32×6, 1 mip, RGBA16F (~48 KB)\n'
            '- specularCube: 256×256×6, 6 mips, RGBA16F (~4 MB)\n'
            '- brdfLut: 256×256, RG16F (~256 KB)\n'
            '运行时只需 3 次纹理采样，基本零开销。'
        ),
    },
    {
        'num': '4',
        'title': '屏幕空间技术：SSAO / GTAO / SSR / SSGI / GTGI',
        'subtitle': '从 G-Buffer + 上一帧 HDR 推导间接光照',
        'intro': (
            '屏幕空间技术利用当前帧的 G-Buffer（深度、法线）和上一帧的 HDR 颜色缓冲来近似局部光照效果。'
            '它们不需要预计算或场景体素化，只依赖当前可见的屏幕内容，'
            '因此成本低、易于集成，但无法获取屏幕外或被遮挡物体的信息。'
            'SomeGI 实现了五类互补的屏幕空间技术。'
        ),
        'principle': (
            '【SSAO — 屏幕空间环境光遮蔽】\n'
            '每个像素在法线半球内采样 16 个点，将其投影到屏幕空间比较深度。'
            '若采样点的场景深度比射线深度更近，则该采样点产生遮挡。AO = 1 - 平均遮挡率。\n'
            '简单高效，但对倾斜表面有"光晕"伪影（采样点穿过表面后方的空白区域被误判为被遮挡）。\n\n'
            '【GTAO — Ground Truth Ambient Occlusion (Jimenez 2016)】\n'
            '替代 SSAO 的更精确方案。沿 4 个切片方向搜索最大水平角（horizon angle），'
            '对每个切片进行解析积分，而非离散计数。完全消除了 SSAO 的光晕问题，'
            '在凹角和褶皱处精度更高。计算成本约为 SSAO 的 1.5–2×。\n\n'
            '【SSR — 屏幕空间反射】\n'
            '对高粗糙度以下的像素（roughness < threshold），沿反射方向在世界空间进行线性 ray march，'
            '每步投影到屏幕空间与深度缓冲比较。命中后采样上一帧 HDR 颜色作为反射色。\n'
            '通过 alpha 通道 fader（屏幕边缘衰减 + 粗糙度衰减）与 IBL specular 平滑混合。\n\n'
            '【SSGI — 屏幕空间全局光照】\n'
            '每像素在法线半球上采样 8 个 cosine 加权方向（Hammersley 序列 + 每像素哈希抖动），'
            '沿方向在世界空间 ray march。命中后采样上一帧 HDR。\n'
            '输出 alpha = hits / sampleCount 作为置信度，直接与 IBL diffuse 混合。\n'
            '时序累积（α=0.92）重投影上一帧的结果到当前帧以降低噪声。\n\n'
            '【GTGI — Ground Truth GI (Sucker Punch 2024 风格)】\n'
            '对 GTAO 的扩展：沿 4 个切片方向搜索水平线时，每当新样本"抬高"水平线，'
            '记录该角度增量 dCos 作为该样本立体角的增量，并从上一帧 HDR 采样其辐射度。\n'
            '结果累加 = Σ radiance · dCos。与 SSGI 相比，噪声结构从散斑变为平滑条纹，'
            '在相同样本数下噪声特性显著更好。'
        ),
        'implementation': (
            '所有屏幕空间技术共享以下资源：\n'
            '- GBuffer: albedo+metallic (RGBA8), normal+roughness (RGBA16F), emissive+AO (RGBA8), depth (D32)\n'
            '- hdrPrev: 上一帧 HDR 颜色 (RGBA16F)\n'
            '- 光照阶段统一进行 5×5 depth-bilateral 滤波平滑所有屏幕空间输出的噪声\n\n'
            'GTAO/GTGI 共享相同的水平线搜索框架（4 slices），与 SSGI 共用输出纹理（互斥）。'
        ),
    },
    {
        'num': '5',
        'title': 'RSM — 反射阴影贴图 (Reflective Shadow Maps)',
        'subtitle': 'Dachsbacher & Stamminger 2005',
        'intro': (
            'RSM 是经典的一次反弹 GI 技术，特别适合受定向光源（太阳）强光影响的户外或大窗户室内场景。'
            '它将太阳视角的阴影贴图升级为"光源视角的 G-Buffer"，'
            '每个 RSM 纹素作为虚拟点光源（VPL）向场景注入间接光照。'
        ),
        'principle': (
            'RSM 的核心思想：从光源视角看，每个可见的表面微元都是被光源直接照亮的"次级光源"，'
            '它可以把接收到的直接光重新发射到周围空间。\n\n'
            '1. 几何阶段（RsmGeometryPass）：从太阳方向渲染场景到 MRT：\n'
            '   - RT0: worldPosition (RGBA16F)\n'
            '   - RT1: worldNormal (RGBA16F)\n'
            '   - RT2: flux = albedo · sunColor · NdotL (RGBA16F)\n'
            '   - Depth: D32_SFLOAT\n\n'
            '2. 采样阶段（RsmSamplePass）：对每个屏幕像素：\n'
            '   - 将 worldPos 投影到 RSM 的 UV 空间\n'
            '   - 在中心 UV 周围用 Hammersley 序列采样 N=64 个圆盘样本\n'
            '   - 对每个 VPL 计算几何贡献：\n'
            '     weight = flux_vpl · max(0, N_vpl·dir) · max(0, N·(-dir)) · offset² / dist⁴\n'
            '   - offset²/d⁴ 中 offset² 补偿均匀圆盘采样，d⁴ 是点到点几何衰减\n'
            '   - alpha = 命中置信度（frustum 覆盖 + sun 可见比例）\n\n'
            '3. 光照集成：rsmF.rgb · albedo · (1-metallic)，按 rsmF.a 与 IBL diffuse 混合。'
        ),
        'implementation': (
            '关键参数：\n'
            '- RSM 分辨率: 512×512\n'
            '- 采样数: N=64 disk 样本/像素\n'
            '- 太阳正交投影范围: 由场景 AABB 自动计算\n'
            '- 不支持多光源 RSM（M5 仅单 sun）\n'
            '- 不支持间接阴影（VPL 之间无遮挡计算）\n\n'
            '局限：只有被太阳看见的表面才能成为 VPL；单次反弹；无 VPL 间接阴影。'
        ),
    },
    {
        'num': '6',
        'title': 'LPV — 光传播体积 (Light Propagation Volumes)',
        'subtitle': 'Kaplanyan / CryEngine 3, 2009',
        'intro': (
            'LPV 将场景离散为粗粒度的 3D 网格（32³），每个格子用一阶球谐函数（SH，4 个系数）'
            '表示该位置的入射辐射度。通过注入（从 RSM）和迭代传播（6-邻居 SH 转移），'
            '光可以在体积中"流动"到 RSM 视锥体之外的区域，实现绕过角落的光传播。'
            'LPV 是体积 GI 中最便宜的选择，但只能捕获低频漫反射。'
        ),
        'principle': (
            '算法分三阶段：\n\n'
            '1. 注入（LpvInjectPass）：\n'
            '对 RSM 中每个非零通量的纹素，将 VPL 的辐射度投影到其所在格子的一阶 SH 中：\n'
            '  sh_k = flux · cosineLobeSh(d_sun) \n'
            'cosineLobeSh 使用 A_l 卷积系数：l=0 → π·Y₀₀, l=1 → (2π/3)·Y₁_m。\n\n'
            '2. 传播（LpvPropagatePass, 8 次迭代）：\n'
            '对每个接收格子，从 6 个邻居的每个邻居通过 5 个面传输光照：\n'
            '  - 直接面（1 个）：立体角 0.4007 sr/π\n'
            '  - 侧面（4 个）：立体角 0.4234 sr/π 每个\n'
            '  - 总计约 2.09 sr/π ≈ 半球 2/3 的覆盖率\n'
            '对每个面，在传输方向上评估邻居的 SH 得到辐射度 L，投影到接收格子的 SH：\n'
            '  dst_sh += L · cosineLobeSh(dir) · solidAngle\n'
            '使用 ping-pong 两套网格迭代 8 次。\n'
            'Geometry Volume (GV) 存储每格遮挡信息，传播时在源-目标中点采样 GV 进行衰减。\n\n'
            '3. 评估（lighting.slang 内联）：\n'
            '世界坐标 → 三线性采样 SH 网格 → evalSh(sh, N) 重建辐照度 → 乘 BRDF。\n'
            '网格边缘用 smoothstep 5% 渐变防止硬边界。'
        ),
        'implementation': (
            '关键参数：\n'
            '- 分辨率: 32³\n'
            '- 数据: 3 张 RGBA16F 3D 纹理 (R/G/B 各 4 个 SH 系数)，ping-pong 两套\n'
            '- 显存: 约 2.25 MB (含 GV)\n'
            '- 传播迭代: 8 次（默认），每次 dispatch 32³ thread\n'
            '- 漏光缓解: GV 遮挡（B.8）在传播中衰减穿过墙体的光\n\n'
            '限制：SH-1 只捕获最低频光照（无锐利阴影/光泽反射）；粗网格导致近处光照块状化；'
            'last-writer-wins 注入在多 VPL 落入同一格子时损失能量。'
        ),
    },
    {
        'num': '7',
        'title': 'VXGI — 体素锥追踪全局光照 (Voxel Cone Tracing GI)',
        'subtitle': 'Crassin et al., Nvidia, 2011',
        'intro': (
            'VXGI 将场景体素化到 128³ 的 3D 网格中，每体素存储反照率和透明度。'
            '通过构建不透明度加权的 mip 链，着色时从接收点沿多个锥方向追踪锥体（cone trace），'
            '利用不同 mip 级别实现自适应的远-粗/近-细采样。'
            '相比 LPV，锥追踪天然带遮挡（不透明度累积），不会漏光；相比 RSM，不受太阳视锥体限制。'
        ),
        'principle': (
            'VXGI 的五个阶段：\n\n'
            '1. 体素化（vxgi_voxelize）：\n'
            '每个三角形采样 7 个点（3 顶点 + 3 边中点 + 重心），将 (albedo, 1) 写入对应体素。'
            '简化了 Crassin 的保守光栅化，但适合大多数场景。\n\n'
            '2. 光照注入（vxgi_inject）：\n'
            '从 RSM 读取每个 VPL 的 worldPos 和 flux，覆写对应体素的 RGB（保留 alpha）。\n\n'
            '3. Mipmap（vxgi_mipmap）：\n'
            '每 2×2×2 体素合并为 1 个：dst.rgb = Σ(src.rgb·src.a) / Σ(src.a)，dst.a = Σ(src.a)/8。\n'
            '不透明度加权的平均避免了空洞或边界体素污染颜色。\n\n'
            '4. 各向异性 Alpha（vxgi_aniso_build, B.6）：\n'
            '在 mip 1+ 级，沿 X/Y/Z 三个轴分别计算 Beer 复合透明度。'
            '锥追踪时使用方向加权的 alpha（max(|dir.x|·ani.r, |dir.y|·ani.g, |dir.z|·ani.b)），大幅减少薄墙漏光。\n\n'
            '5. 锥追踪（lighting.slang 内联）：\n'
            '6 个锥方向（N 半球的 cosine 加权 cluster）从表面出发，沿方向前进 24 步。\n'
            '每步根据 cone 直径选择 mip 级别: mip = log₂(diameter / cellSize)。\n'
            '前向合成：acc += trans · voxel.rgb · voxel.a; trans *= (1 - voxel.a)。\n'
            'trans < 0.01 时提前终止。最终按 cone 权重归一化。\n\n'
            '6. 多次反弹 Relight（vxgi_relight, C.2）：\n'
            '可选地在体素网格自身做一次锥追踪（6 主轴方向），将收集到的辐射度加回每个体素。'
            '模拟场景内部的光线多次反射。'
        ),
        'implementation': (
            '关键参数：\n'
            '- 体素分辨率: 128³\n'
            '- Mip 链: 8 级 (log₂128 + 1)\n'
            '- 格式: RGBA16F, 含 mip 链约 18.3 MB\n'
            '- 各向异性 alpha: 额外 ~18.3 MB\n'
            '- 6 个锥, 24 步/锥, 每像素 144 次 3D 纹理采样（主要开销）\n'
            '- 每帧完整重建（voxelize + inject + mipmap + 可选 aniso/relight）\n\n'
            '限制：计算散播体素化（非保守光栅）可能遗漏细小几何体；单 cascade 无远距离覆盖；'
            '每帧重建对动态场景昂贵。'
        ),
    },
    {
        'num': '8',
        'title': 'PRT — 预计算辐射度传输 (Precomputed Radiance Transfer)',
        'subtitle': 'Sloan et al., SIGGRAPH 2002',
        'intro': (
            'PRT 将光照分为两个独立的 SH 投影部分：光源分布（light SH）和每点可见性传输函数（transfer SH）。'
            '传输函数在场景加载时一次性烘焙，运行时只需将光源投影到 SH 并与传输 SH 做点积，'
            '即可得到带遮挡的漫反射间接光照。PRT 的独特优势在于传输与光源无关——'
            '太阳方向 / 强度变化时无需重新烘焙。'
        ),
        'principle': (
            'PRT 的核心分解：\n\n'
            '  E(x, N) = ∫ L(d) · max(0, N·d) · V(x, d) dω\n'
            '          ≈ max(0, Σ_k l_k · v_k(x) · A_l(k) · Y_k(N))\n\n'
            '其中：\n'
            '- L(d) = 光源 radiance → l_k = light SH 系数（CPU 每帧投影 sun 方向）\n'
            '- V(x, d) = 二元可见性 → v_k(x) = transfer SH 系数（一次烘焙）\n'
            '- A_l(k) = Ramamoorthi/Hanrahan clamped-cosine 卷积系数：\n'
            '  l=0 → π, l=1 → 2π/3, l=2 → π/4, l=3 → 0\n\n'
            '烘焙过程（prt_bake）：\n'
            '对 32³ 网格中的每个体素中心：\n'
            '1. 在球面上生成 64 个 Hammersley 均匀样本\n'
            '2. 对每方向通过体素网格 ray march 确定可见性 V ∈ [0, 1]\n'
            '3. 累加: v_k(x) = (4π/N) · Σ V(d_i) · Y_k(d_i)\n'
            '4. 输出 transfer SH 系数到 RGBA16F 3D 纹理\n\n'
            '运行时重建（lighting.slang）：\n'
            '1. CPU 每帧投影 sun 到 SH: l_k = sunIntensity · sunColor · Y_k(sunDir)\n'
            '2. GPU: 三线性采样 transfer SH，逐元素乘 light SH 和 A_l\n'
            '3. 通过 Y_k(N) 重建辐照度 → irr · albedo · (1-metallic) / π'
        ),
        'implementation': (
            '关键参数：\n'
            '- 体素分辨率: 32³\n'
            '- SH 阶数: 支持 SH4 (4 系数, 1 张纹理)、SH9 (9 系数, 3 张)、SH16 (16 系数, 5 张)\n'
            '- 显存: SH4 ~2 MB, SH9 ~6 MB, SH16 ~10 MB\n'
            '- 烘焙时间: 一次 ~数百 ms（32³ × 64 ray × 24 steps）\n'
            '- 运行时开销: 1 次 3D 纹理采样 + 若干 dot 运算（基本免费）\n\n'
            '限制：只处理静态几何（动态场景需重新烘焙）；SH-4/9 仅捕获低频光照；无多次反弹。'
        ),
    },
    {
        'num': '9',
        'title': 'DDGI — 动态漫反射全局光照 (Dynamic Diffuse GI)',
        'subtitle': 'Majercik et al., NVIDIA RTXGI, 2020',
        'intro': (
            'DDGI 在场景中布设 3D 探针网格（256 个），每个探针存储全方位的辐照度（八面体编码）'
            '和深度统计数据（Chebyshev 测试用）。每帧每个探针发射 64 条球面斐波那契分布光线，'
            '通过体素网格追踪命中，将结果按方向加权累积到 atlas 纹理中，以极小滞后系数（α=0.05）'
            '进行时序混合。着色时对目标点周围 8 个探针进行三线性插值 + Chebyshev 可见性测试。'
        ),
        'principle': (
            'DDGI 三阶段流程：\n\n'
            '1. 光线追踪更新（ddgi_update）：\n'
            '- 每探针发射 N=64 条光线，方向用球面斐波那契序列均匀分布\n'
            '- 每帧对方向施加随机 Y 轴旋转（时间抖动，确保长时间覆盖均匀）\n'
            '- 在体素网格中 ray march（最多 32 步），返回命中辐射度和距离\n\n'
            '2. Atlas 混合（ddgi_blend）：\n'
            '- 辐照度 atlas：每纹素对应一个方向，对 64 条光线按 cos(texelDir, rayDir) 加权累加，\n'
            '  除以总权重得到新辐照度。与上一帧以 α=0.05 混合（时域平滑）\n'
            '- 距离 atlas：同辐照度，但加权余弦幂提高到 50（更锐利的方向选择性）\n'
            '  存储 (mean, mean²) 用于 Chebyshev 不等式计算方差\n\n'
            '3. 探针分类（ddgi_classify）：\n'
            '- 若大部分光线命中距离 < closeHitDist，探针标记为 inactive（在墙体内）\n\n'
            '4. 着色采样（lighting.slang 内联）：\n'
            '- 对 worldPos 周围 8 个探针进行三线性插值\n'
            '- 背面剔除：探针朝像素方向与 N 同向才加权重\n'
            '- Chebyshev 可见性: vis = σ² / (σ² + (d - μ)²)，立方后增强对比\n'
            '- 在法线 N 方向采样辐照度 atlas → irr · albedo · (1-metallic) / π'
        ),
        'implementation': (
            '关键参数：\n'
            '- 探针网格: 8×4×8 = 256 探针（按场景 AABB 摆放，5% padding）\n'
            '- 辐照度 atlas: 8² 纹素/探针, 八面体映射, RGBA16F\n'
            '- 距离 atlas: 16² 纹素/探针, RG16F\n'
            '- 光线: 64/探针, 共 16,384 条/帧\n'
            '- 时序混合: α=0.05（95% 历史, ~140 帧时域窗口）\n'
            '- 依赖 VXGI 体素网格做光线追踪的加速结构\n\n'
            '限制：探针密度决定空间分辨率；时序滞后在快速光照变化时可见；Chebyshev 可见性偏软。'
        ),
    },
    {
        'num': '10',
        'title': 'SDFGI — 有符号距离场全局光照 (Signed Distance Field GI)',
        'subtitle': 'Godot 4 JFA + Sphere Tracing',
        'intro': (
            'SDFGI 是 Godot 4 引擎的主 GI 方案。SomeGI 的实现将体素网格转换为无符号距离场（UDF），'
            '利用 Jump Flood Algorithm (JFA) 高效计算最近占用体素的距离，然后使用 sphere tracing '
            '进行精确的光线追踪。相比 VXGI 的固定步长锥追踪，sphere tracing 能自适应步进，'
            '对薄几何体更精确且步数更少。'
        ),
        'principle': (
            'SDFGI 四个阶段：\n\n'
            '1. 种子生成（sdfgi_seed）：\n'
            '对 128³ 体素网格每格检查 alpha ≥ threshold，占用的体素标记为"种子"（存储自身坐标）。\n\n'
            '2. JFA 传播（sdfgi_jfa, 7 次迭代）：\n'
            '跳转洪水算法：第 k 次迭代检查步长 s = 2^(6-k) 的 27 个邻居（3×3×3），'
            '记录距当前格子最近的种子位置。\n'
            '7 次迭代（k=64,32,16,8,4,2,1）后每个格子都知道离自己最近的占用体素。\n'
            'ping-pong 两套 128³ RGBA16F 缓冲交换（共 32 MB 临时存储）。\n\n'
            '3. 最终化（sdfgi_finalize）：\n'
            '将 JFA 结果转为标量 UDF: d = length(cellPos - nearestSeedPos) in cell units。\n'
            '存储为 R16F 3D 纹理（4 MB）。\n\n'
            '4. Sphere Tracing（sdfgi_trace）：\n'
            '每像素沿 N 个 cosine 加权半球方向进行 sphere tracing：\n'
            '  while traveled < maxDist:\n'
            '    d = sdfTex.Sample(uv)\n'
            '    if d < hitEps: hit! 采样体素 radiance\n'
            '    step = max(d, 0.5)  // 最小步长防止卡住\n'
            '    p += dir * step\n'
            'Sphere tracing 保证不进入占用体素，且步长自适应于局部几何密度。'
        ),
        'implementation': (
            '关键参数：\n'
            '- UDF 分辨率: 128³ R16F = 4 MB\n'
            '- JFA 临时存储: 2 × 128³ RGBA16F = 32 MB\n'
            '- JFA 迭代: 7 次（log₂128）\n'
            '- Sphere tracing: 4-16 rays × 48 max steps\n'
            '- 依赖 VXGI 体素网格作为占用信息源\n\n'
            '限制：无符号距离（不区分内外）只支持从外部向场景追踪；单 cascade 无远距离覆盖。'
        ),
    },
    {
        'num': '11',
        'title': 'ReSTIR DI — 储层时空重要性重采样直接光照',
        'subtitle': 'Bitterli et al., SIGGRAPH 2020 / Wyman et al., HPG 2021',
        'intro': (
            'ReSTIR 本身不是 GI 技术，而是一种高效的直接光照采样框架，解决"多光源场景中如何高效选择重要光源"的问题。'
            '核心思想：每个像素维护一个储层（reservoir），通过加权储层采样（RIS）从大量候选光源中概率性地选出重要光源，'
            '然后通过空间复用让邻居像素共享采样结果。最终每个像素只需 1 次可见性测试，'
            '而不是为每个候选光源都做可见性测试。'
        ),
        'principle': (
            'ReSTIR DI 三阶段：\n\n'
            '1. 初始采样 RIS（restir_init）：\n'
            '每个像素从 M=8 个均匀随机光源中选取，计算目标 PDF p_hat (BRDF · Le · cos / dist²)。\n'
            '通过加权储层采样选出一个光源，储层记录：selectedLight, wSum, M, W。\n'
            '最终无偏权重 W = wSum / (M · p_hat(selected))。此阶段无可见性测试。\n\n'
            '2. 空间复用（restir_spatial）：\n'
            '每个像素从 K=4 个屏幕空间邻居读取储层，用深度/法线阈值拒绝跨表面的邻居。\n'
            '对每个邻居的选中光源在当前像素重新计算 targetPDF，通过 balance heuristic 合并储层。\n'
            '合并后有效采样数 ≈ M × (K+1) = 40。仍无可见性测试。\n\n'
            '3. 着色（restir_shade）：\n'
            '解包储层得到唯一选中的光源，通过体素网格 ray march 执行 1 次可见性测试。\n'
            '完整光照贡献 = (albedo/π) · Li · NdotL / dist² · vis · W · intensityScale。\n'
            '输出到 RGBA16F 屏幕纹理，lighting.slang 直接加到 hdrColor。\n\n'
            'ReSTIR 的核心收益：每个像素只需 1 次可见性测试即可等效评估数十个光源，'
            '将直接光照的成本从 O(N_lights × visibility) 降为 O(1 × visibility)。'
        ),
        'implementation': (
            '关键参数：\n'
            '- 储层: 2 × RGBA32UI 屏幕分辨率 (16 MB @1080p)\n'
            '- 候选数 M: 8\n'
            '- 空间邻居 K: 4\n'
            '- 可见性: 体素网格 ray march（复用 VXGI 网格）\n'
            '- 演示光源: 8 个点光源（场景 AABB 角落）\n\n'
            '限制：当前仅时序复用尚未实现（D.2）；仅支持点光源（D.6 计划支持 emissive triangle）。'
        ),
    },
]

for tech in gi_techniques:
    doc.add_heading(f'{tech["num"]}. {tech["title"]}', level=1)
    doc.add_paragraph(tech['subtitle'], style='Normal')

    doc.add_heading('概述', level=2)
    doc.add_paragraph(tech['intro'])

    doc.add_heading('基本原理', level=2)
    doc.add_paragraph(tech['principle'])

    doc.add_heading('实现细节', level=2)
    doc.add_paragraph(tech['implementation'])

    doc.add_page_break()

# ── 12. 性能与效果对比 ──────────────────────────────────
doc.add_heading('12. 性能与效果对比', level=1)

doc.add_heading('12.1 各 GI 技术特性对比', level=2)

add_table(
    ['技术', '类型', '反弹次数', '遮挡感知', '动态几何', '镜面反射', '分辨率依赖', '预计算'],
    [
        ['IBL', '环境探针', '∞ (预计算)', '否', '是', '是', '否', '是（启动时）'],
        ['SSAO/GTAO', '屏幕空间', 'N/A (遮蔽)', '部分', '是', '否', '是', '否'],
        ['SSR', '屏幕空间', '1', '否', '是', '是', '是', '否'],
        ['SSGI', '屏幕空间', '1+feedback', '否', '是', '否', '是', '否'],
        ['GTGI', '屏幕空间', '1+feedback', '部分', '是', '否', '是', '否'],
        ['RSM', '体积/空间', '1', '否', '是', '否', '否', '否'],
        ['LPV', '体积/空间', '多次', '部分 (GV)', '是', '否', '否', '否'],
        ['VXGI', '体积/空间', '1+relight', '是 (cone)', '是', '否', '否', '否'],
        ['PRT', '预计算', '1', '是 (baked)', '否', '否', '否', '是（场景加载）'],
        ['DDGI', '探针/空间', '多次', '是 (Chebyshev)', '是', '否', '否', '否'],
        ['SDFGI', '体积/空间', '1', '是 (SDF)', '是', '否', '否', '否'],
        ['ReSTIR DI', '采样框架', 'N/A (直接光)', '是 (1 次)', '是', '是 (specular)', '否', '否'],
    ],
    [2.5, 2, 2, 2, 2, 2, 2.5, 2.5]
)

doc.add_heading('12.2 显存占用估算', level=2)
doc.add_paragraph('以下为各技术附加显存（不含共享的 G-Buffer/场景数据），基于 1080p 分辨率计算：')

add_table(
    ['技术', '主要资源', '估算显存', '备注'],
    [
        ['IBL', 'Cubemap × 3 + LUT', '~14 MB', '启动时一次性分配'],
        ['SSAO/GTAO', '输出纹理 (R8)', '~2 MB', '与 GTAO 共享'],
        ['SSR', '输出纹理 (RGBA16F)', '~16 MB', ''],
        ['SSGI/GTGI', '输出 + History (RGBA16F)', '~32 MB', ''],
        ['RSM', '4 张 RSM 纹理 (512²)', '~12 MB', '含 position/normal/flux/depth'],
        ['LPV', '3D 纹理 × 6 + GV', '~2.25 MB', 'ping-pong 两套 32³'],
        ['VXGI', '3D 纹理 + mip 链 + aniso', '~37 MB', '128³ RGBA16F'],
        ['PRT', '3D 纹理 × 1~5', '2~10 MB', '取决于 SH 阶数'],
        ['DDGI', 'Atlas × 2 + rayBuf + states', '~3 MB', '256 探针'],
        ['SDFGI', 'UDF + JFA 临时 × 2', '~36 MB', '含临时存储'],
        ['ReSTIR DI', '储层 (RGBA32UI × 2)', '~32 MB', ''],
    ],
    [2.5, 4, 2.5, 5]
)

doc.add_heading('12.3 画质特性对比', level=2)

add_table(
    ['技术', '色溢效果', '漏光', '噪声水平', '适用场景'],
    [
        ['IBL', '无局部色溢', '无', '无噪声', '所有场景（回退方案）'],
        ['SSGI', '较弱', '屏幕外无', '中等（需滤波）', '近景补充'],
        ['GTGI', '弱至中等', '屏幕外无', '较低', '近景补充'],
        ['RSM', '强（定向光方向）', '少', '中等', '户外 / 大窗户室内'],
        ['LPV', '中等', '较多（薄墙）', '低', '低频室内环境'],
        ['VXGI', '中等至强', '少（aniso 后极少）', '低', '所有场景（主力）'],
        ['PRT', '中等', '无（baked vis）', '无噪声', '静态场景 + 动态光源'],
        ['DDGI', '中等至强', '少（Chebyshev）', '低', '所有场景'],
        ['SDFGI', '中等', '极少', '中等（需时序）', '薄几何体场景'],
        ['ReSTIR DI', 'N/A (直接光)', 'N/A', '处理得当后低', '多光源场景'],
    ],
    [2, 2.5, 2, 2, 5.5]
)

doc.add_heading('12.4 GPU 管线阶段开销分布', level=2)
doc.add_paragraph(
    '以下为 Sponza 场景在 Intel UHD 770 上的典型帧时间分布（1080p, 60 FPS 上限, GPU 约 4-5 ms/帧）：\n\n'
    '• G-Buffer 渲染: ~1.2 ms\n'
    '• AO+SSR+SSGI 屏幕空间: ~0.5 ms (合计)\n'
    '• VXGI 全管线: ~2.5 ms (voxelize+inject+mipmap+cone trace)\n'
    '• DDGI: ~1 ms (256 探针 × 64 光线)\n'
    '• LPV: ~0.3 ms\n'
    '• Lighting: ~0.8 ms\n'
    '• Tonemap+ImGui: ~0.2 ms\n\n'
    '注：并非所有 GI 都同时激活。典型模式（VXGI 或 DDGI 单独运行）总计 GPU 约 4-5 ms。'
    '各技术的开销排序（从低到高）：\n'
    'IBL < SSGI/GTAO/SSR < RSM < LPV < DDGI < SDFGI < VXGI'
)

doc.add_page_break()

# ── 13. 总结 ─────────────────────────────────────────────
doc.add_heading('13. 总结与展望', level=1)

doc.add_heading('13.1 各技术适用场景总结', level=2)
doc.add_paragraph(
    '• IBL: 所有场景的基础环境光，提供方向感的环境反射和漫反射\n'
    '• 屏幕空间技术 (SSAO/SSR/SSGI/GTGI): 轻量级补充，零预计算，适合近景细节\n'
    '• RSM: 最适合户外 / 太阳直射场景的一次反弹\n'
    '• LPV: 最便宜的体积 GI，适合低频室内环境光传播\n'
    '• VXGI: 主力体积 GI，对遮挡精确且支持多次反弹，适合所有场景\n'
    '• PRT: 静态场景 + 动态光源的理想组合，运行时几乎零开销\n'
    '• DDGI: 时间平滑的探针 GI，适合缓慢变化的间接光\n'
    '• SDFGI: 更精确的光线追踪（替代 VXGI 锥追踪），适合薄几何体场景\n'
    '• ReSTIR DI: 高效处理大量动态点光源的直接光照'
)

doc.add_heading('13.2 后续计划', level=2)
doc.add_paragraph(
    '短期（高 ROI）：\n'
    '• D.1 SDFGI 时序累积 — 加入重投影 + α-blend 降低噪点\n'
    '• D.2 ReSTIR DI 时序复用 — 实现论文核心的 temporal reservoir reuse\n'
    '• D.5 VXGI 6-cone hemisphere — 替换当前 6 主轴粗略近似\n\n'
    '中期（放大器）：\n'
    '• D.3 GI temporal denoiser (SVGF 风格) — 所有 GI 通道受益\n'
    '• E.3 TAA 抗锯齿 — 与 denoiser 共享 history\n\n'
    '长期：\n'
    '• L 阶段 Lumen-lite — 复用现有基础设施复刻 UE5 Lumen 的 80% 视觉效果\n'
    '• F 阶段硬件 RT — 等换支持 VK_KHR_ray_tracing 的 GPU 后启用'
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 报告完 —')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), 'SomeGI_全局光照技术实现报告.docx')
doc.save(output_path)
print(f'报告已保存到: {output_path}')
